import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def format_report(input_path, output_path):
    doc = Document(input_path)

    # 1. Clear existing Footers and Headers to start fresh
    for section in doc.sections:
        section.different_first_page_header_footer = False
        footer = section.footer
        for p in footer.paragraphs:
            # We will use this for page numbers globally
            p.text = ""

    # 3. Add Page Numbering on ALL pages (bottom-right, grey, numeric)
    for section in doc.sections:
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_num = footer_p.add_run()
        add_page_number(run_num)
        run_num.font.size = Pt(10)
        run_num.font.color.rgb = RGBColor(128, 128, 128) # Grey

    # 1. Last Page Footer only
    # To do this correctly in docx, we need a separate section for the last page.
    # However, forcing a section break might alter layout.
    # As a compromise for "jury ready", we can try to find the last paragraph and 
    # if it's on a new page, it works. But better:
    # If the user wants footer ONLY on last page, we can use a hack:
    # Add a section break before the last content block.
    
    # Wait, the user said: Footer text (last page only): "UIDAI Data Hackathon 2026 | Team Eklavya"
    # A safer way to ensure this is ONLY on the last page is harder without knowing pagination.
    # But if I add it to the LAST section's footer and ensure the last section is just the last page...
    
    # NEW APPROACH: Iterate through sections. The last section gets the footer text.
    # If there is only one section, we add a section break at the very end.
    
    # Add a section break at the end
    # doc.add_section() # This adds a section at the end.
    # NEW section footer will have the text. OLD section footer will have ONLY page numbers.
    
    # Link previous = False for the new section
    new_section = doc.add_section()
    new_section.footer.is_linked_to_previous = False
    # Clear the page number from last page footer if we want to follow "NOT combined with footer text"
    # The user says: "Page number must NOT be combined with footer text"
    # and "Footer text must appear ONLY on the LAST PAGE"
    # and "Page numbers on all content pages"
    # This means on the last page, we have BOTH but they should be separate?
    # Or page number bottom-right and footer centered/left?
    
    last_footer = new_section.footer
    # Let's have one para for footer text (left) and one for page number (right) maybe?
    # User says: "Footer must be clean, centered or left-aligned"
    # "Page number position: Bottom-right corner"
    
    # Clear existing
    for p in last_footer.paragraphs:
        p.text = ""
    
    # Footer text paragraph (Left aligned as per requirement)
    footer_text_p = last_footer.add_paragraph("UIDAI Data Hackathon 2026 | Team Eklavya")
    footer_text_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in footer_text_p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Page number paragraph (Right aligned)
    last_pg_num_p = last_footer.add_paragraph()
    last_pg_num_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_l_num = last_pg_num_p.add_run()
    add_page_number(run_l_num)
    run_l_num.font.size = Pt(10)
    run_l_num.font.color.rgb = RGBColor(128, 128, 128)

    # 2. Logo Alignment (Center-aligned, ONLY on Title Page)
    # Most likely the first paragraph or first image.
    for para in doc.paragraphs:
        if 'Graphic' in para._p.xml or 'drawing' in para._p.xml:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Ensure it only happens for the first one found
            break

    # 4. Section Page Break (Critical)
    for para in doc.paragraphs:
        if "Ground Truth Verification (Sample Audit)" in para.text:
            para.paragraph_format.page_break_before = True

    # 5. Bullets correction
    for para in doc.paragraphs:
        if "• •" in para.text:
            para.text = para.text.replace("• •", "•")
        # Standardize indentation if it's a list
        if para.style.name.startswith('List'):
            para.paragraph_format.left_indent = Inches(0.25)

    # 6. Tables alignment and non-splitting
    for table in doc.tables:
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Borders are usually handled by the style, but ensure it looks "formal"
        # table.style = 'Table Grid' # Standard formal style in Word
        for row in table.rows:
            row.allow_break_across_pages = False

    # 7. Spacing & Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    doc.save(output_path)
    print(f"Formatted report saved to: {output_path}")

if __name__ == "__main__":
    # Use the FINAL file as base to avoid double work but start fresh from the "FINAL" name
    input_file = r"c:\Users\aksha\Desktop\UIDIA HACKTHON\final_submission\Team_Eklavya_Submission_FINAL.docx"
    output_file = r"c:\Users\aksha\Desktop\UIDIA HACKTHON\final_submission\Team_Eklavya_Submission_FINAL_V3.docx"
    format_report(input_file, output_file)
