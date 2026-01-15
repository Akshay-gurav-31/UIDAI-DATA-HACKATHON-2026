import docx
import os

def audit_formatting(path):
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return

    doc = docx.Document(path)
    print(f"--- Formatting Audit for: {os.path.basename(path)} ---")
    
    issues = []
    
    # Check for excessive empty paragraphs (Gap analysis)
    consecutive_empty = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            consecutive_empty += 1
        else:
            if consecutive_empty >= 2:
                issues.append(f"[GAP] Excessive spacing ({consecutive_empty} lines) before: \"{text[:40]}...\"")
            consecutive_empty = 0
            
            # Check for Orphan Headings (Heading at the end of a section/page usually followed by empty)
            if p.style.name.startswith('Heading'):
                if i + 1 < len(doc.paragraphs):
                    next_p = doc.paragraphs[i+1].text.strip()
                    if not next_p:
                        issues.append(f"[LAYOUT] Potential Orphan Heading: \"{text}\" is followed by a blank line.")

    # Check Table Alignment
    for i, table in enumerate(doc.tables):
        alignment = getattr(table, 'alignment', 'None')
        if alignment is None:
            issues.append(f"[TABLE] Table {i+1} has no explicit alignment (might be off-center).")
            
    # Check for Double Bullets (Sanity Check)
    for p in doc.paragraphs:
        if "• •" in p.text:
            issues.append(f"[BULLET] Double bullet found in paragraph: \"{p.text[:30]}...\"")

    if not issues:
        print("No mechanical formatting mistakes found.")
    else:
        for issue in issues:
            print(issue)

if __name__ == "__main__":
    audit_formatting(r"c:\Users\aksha\Desktop\UIDIA HACKTHON\final_submission\Team_Eklavya_Submission_FINAL.docx")
