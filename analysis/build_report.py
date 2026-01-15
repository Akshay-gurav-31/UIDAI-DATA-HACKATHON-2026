from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_DIR = r"c:/Users/aksha/Desktop/UIDIA HACKTHON"
IMAGE_DIR = os.path.join(BASE_DIR, "final_submission", "images")
LOGO_PATH = os.path.join(IMAGE_DIR, "TEAM-EKLAVYA-logo.png")

def create_submission():
    doc = Document()

    # Style Setup
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Logo and Title
    if os.path.exists(LOGO_PATH):
        doc.add_picture(LOGO_PATH, width=Inches(1.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_heading('UIDAI DATA HACKATHON 2026 - SUBMISSION REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Team Name: Team Eklavya').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph('Date: January 15, 2026').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Live Dashboard Link (High Priority)
    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = link_p.add_run("🔴 LIVE PROTOTYPE: https://uidia-dashboard.vercel.app/")
    run.bold = True
    run.font.color.rgb = None # Default color (or make it blue if library allows simple RGB)
    run.font.size = Pt(12)
    
    # Dashboard Screenshot (The "Hero" Image)
    dashboard_img = os.path.join(IMAGE_DIR, "dashboard.png")
    if os.path.exists(dashboard_img):
        doc.add_picture(dashboard_img, width=Inches(6.5)) # Widescreen impact
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 1: Eklavya Command Center - Real-Time Anomaly Detection Interface", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph() # Spacing

    # 1. Problem Statement and Approach
    doc.add_heading('1. Problem Statement and Approach', level=1)
    doc.add_paragraph(
        "Our team identified monitoring gaps in current UIDAI data pipelines where enrolment flows and update flows "
        "show synchronisation challenges. Our approach focuses on 'Engineering Discipline'—"
        "using microscopic data auditing to identify operational latencies, naming inconsistencies, and reporting cycle patterns "
        "that require administrative attention."
    )
    
    # 2. Datasets Used
    doc.add_heading('2. Datasets Used', level=1)
    datasets = [
        "api_data_aadhar_enrolment (New ID Generation Stream)",
        "api_data_aadhar_demographic (Identity Maintenance Stream)",
        "api_data_aadhar_biometric (Security & Lifecycle Stream)"
    ]
    for ds in datasets:
        doc.add_paragraph(ds, style='List Bullet')

    # 3. Methodology
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph(
        "Our methodology involves three key technical phases:\n"
        "1. Cross-API Delta Auditing: Comparing enrolment headers against update headers to find naming mismatches.\n"
        "2. Temporal Pulse Analysis: Resampling transaction data to identifying artificial reporting cycles.\n"
        "3. Statistical Correlation Modeling: Using Pearson Correlation coefficients to distinguish organic demand from forced administrative batching."
    )

    # 4. Data Analysis and Visualisation
    doc.add_heading('4. Data Analysis and Visualisation', level=1)

    # Finding A
    doc.add_heading('Finding A: The Naming Paradox (Ghost Districts)', level=2)
    doc.add_paragraph(
        "Audit revealed districts like 'Bengaluru Urban' showing massive enrolments (9,340) but zero updates. "
        "Simultaneously, 'Bengaluru South' showed high updates but zero enrolments. "
        "Conclusion: A critical naming mismatch in the backend API headers creates a 'Ghost District' illusion."
    )
    img_a = os.path.join(IMAGE_DIR, "naming_trap_v2.png")
    if os.path.exists(img_a):
        doc.add_picture(img_a, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 2: Cross-API Naming Mismatch (Bangalore Case Study)", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Finding B
    doc.add_heading('Finding B: The Monthly Pulse (Operational Latency)', level=2)
    doc.add_paragraph(
        "Temporal analysis shows 91% of all Aadhaar transactions are reported on the 1st of every month, "
        "with near-zero activity on other days. This proves the system operates as a 'Batch Processor' "
        "rather than a real-time stream, introducing a significant 30-day decision-making latency."
    )
    img_b = os.path.join(IMAGE_DIR, "system_pulse_v2.png")
    if os.path.exists(img_b):
        doc.add_picture(img_b, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 3: Daily Transaction Volume Audit (Evidence of Batching)", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Finding C
    doc.add_heading('Finding C: Update Load Synchronisation Pattern', level=2)
    doc.add_paragraph(
        "We discovered a 0.99 (99%) mathematical correlation between child and adult updates. "
        "This high correlation indicates synchronized processing cycles where adult and child update streams "
        "show coupled operational patterns, suggesting opportunities for load distribution optimization."
    )
    doc.add_paragraph(
        "Note: High correlation indicates synchronized processing cycles, not individual update behavior. "
        "This reflects aggregated administrative patterns in the reporting infrastructure.",
        style='Intense Quote'
    )
    img_c = os.path.join(IMAGE_DIR, "adult_tsunami_v2.png")
    if os.path.exists(img_c):
        doc.add_picture(img_c, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 4: Mathematical Correlation Analysis (Synchronized Update Patterns)", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5. Solution / Recommendations
    doc.add_heading('5. Solution / Recommendations', level=1)
    solutions = [
        "Naming Standardization Monitoring: Automated data quality signals to flag naming inconsistencies across APIs.",
        "Enhanced Reporting Cadence: Transition high-growth administrative units from monthly to daily ingestion cycles.",
        "Segmented Load Balancing: Optimize processing workflows to distribute update streams more efficiently."
    ]
    for sol in solutions:
        doc.add_paragraph(sol, style='List Bullet')

    # 6. Conclusion / Impact
    doc.add_heading('6. Conclusion / Impact', level=1)
    doc.add_paragraph(
        "By addressing these monitoring gaps, administrative units can benefit from enhanced reporting cadence, "
        "improved data quality signals, and optimized processing workflows. This supports a more responsive and reliable Aadhaar ecosystem."
    )
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("Disclaimer: ").bold = True
    disclaimer.add_run(
        "This prototype demonstrates how aggregated Aadhaar datasets can be operationalized for continuous system health monitoring. "
        "All findings are based on publicly available synthetic/sample datasets and do not represent official UIDAI operational metrics."
    )
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 7. Jury Pitch
    doc.add_heading('Strategic Pitch: Key Takeaways', level=1)
    pitch = doc.add_paragraph()
    pitch.add_run("Our project provides an 'X-Ray' of UIDAI's data flow. ").bold = True
    pitch.add_run(
        "Team Eklavya used data hygiene and engineering discipline to find three killer "
        "insights that directly impact national-level policy and system reliability."
    )

    # Save
    save_path = os.path.join(BASE_DIR, "final_submission", "eklavya_submission.docx")
    doc.save(save_path)
    print(f"SUCCESS: Final Document saved at {save_path}")

if __name__ == "__main__":
    create_submission()
